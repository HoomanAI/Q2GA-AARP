"""Build a Word document ("Q2GA Detailed Flowcharts and Pseudocode.docx")
that explains the Q2GA algorithm's internal components in more detail than
the user's hand-drawn overview flowchart, and includes pseudocode for the
GA, SA, QGA, and Q2GA algorithms used in the comparative study.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FLOW_DIR = os.path.join(ROOT, "figures", "flowcharts")
OUT_PATH = os.path.join(os.path.dirname(ROOT), "Q2GA Detailed Flowcharts and Pseudocode.docx")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, n, filename, title, description, width=6.0):
    add_heading(doc, f"Figure {n}: {title}", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(FLOW_DIR, filename), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure {n}. {title}. (file: {filename})")
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    desc = doc.add_paragraph(description)
    return desc


def add_pseudocode(doc, title, lines):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr()
    # shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    cell._tc.get_or_add_tcPr().append(shd)

    cell.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def main():
    doc = Document()

    title = doc.add_heading("Q2GA: Detailed Component Flowcharts and Algorithm Pseudocode", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    intro = doc.add_paragraph(
        "This document complements the high-level Q2GA flowchart (Begin -> Generate Initial "
        "Population Q(t) -> Make P(t) -> Evaluate P(t) -> ... -> Update Q(t) with q-gate U(t) -> "
        "Save best individual -> loop) by zooming into each of its internal components: the "
        "quantum chromosome representation and population collapse, the reinforcement-learning "
        "(RL) observation state, the epsilon-greedy action-selection mechanism, the adaptive "
        "quantum rotation gate U(t), and the reward/Q-table update. A side-by-side comparison "
        "of the QGA (fixed-angle) and Q2GA (RL-adaptive) rotation gates is also included. "
        "Pseudocode for all four algorithms used in the comparative study (GA, SA, QGA, Q2GA) "
        "is provided in the final section."
    )

    add_heading(doc, "1. Detailed Component Flowcharts", level=1)

    add_figure(
        doc, 1, "q2ga_master_flow.png",
        "Master diagram -- relating the detailed figures to one Q2GA generation",
        "This diagram shows where each of the following detailed figures fits within a single "
        "generation t of the Q2GA loop shown in the user's overview chart. Stage 1 (Figure 2) "
        "covers the collapse of the quantum population Q(t) into the real-valued population "
        "P(t) and its evaluation. Stage 2 (Figure 3) covers how each individual's RL "
        "observation state O_t,i = (g, h, f) is constructed. Stage 3 (Figure 4) covers how the "
        "RL agent selects an action (operator) using an epsilon-greedy policy over its Q-table. "
        "Stage 4 (Figure 5) covers how that action is translated into an adaptive quantum "
        "rotation gate U(t) applied to selected qubits. Stage 5 (Figure 6) covers reward "
        "computation, the Q-table update, and best-solution tracking, after which the loop "
        "advances to generation t+1. Stages 2-5 correspond to the orange RL-specific boxes in "
        "the original hand-drawn chart (Reward Calculation, Update State, Softmax/epsilon-"
        "greedy, Action Selection, Update Q(t) with q-gate U(t))."
    )

    add_figure(
        doc, 2, "q2ga_chromosome_collapse.png",
        "Quantum chromosome representation and population collapse",
        "Each individual i in the quantum population Q(t) is represented as a string of qubit "
        "pairs (alpha_{i,j}, beta_{i,j}) for j = 1..n_qubits, with alpha^2 + beta^2 = 1, where "
        "|beta_{i,j}|^2 gives the probability of observing bit value 1 for qubit j. To form the "
        "real-coded population P(t) (\"Make P(t) by observing Q(t)\"), each qubit is measured "
        "independently to produce a binary string b_i, which is decoded via bits_to_real() into "
        "a real-valued AARP chromosome (vehicle-route / sequencing genes in [0,1]). This "
        "chromosome is decoded into vehicle routes and a dispatch simulation is run to compute "
        "the multi-objective fitness (weighted travel cost, response-time penalties, and "
        "infeasibility penalties). If an individual's fitness improves on the best solution "
        "found so far, the best-so-far bit string and fitness are updated -- this best-so-far "
        "record is later used by the rotation gate (Figures 5 and 7) to steer qubits toward "
        "promising bit values."
    )

    add_figure(
        doc, 3, "q2ga_observation_state.png",
        "Construction of the RL observation state O_t,i = (g, h, f)",
        "For every individual i, the RL agent observes a discrete state O_t,i = (g, h, f) used "
        "to index its Q-table. The improvement flag g in {0,1} compares the population's mean "
        "fitness at generation t against the running average of the previous GEN_WINDOW = 5 "
        "generations: g = 1 if the population is currently improving, 0 otherwise. The "
        "diversity bin h in {0,1,2,3} is the quartile bin of individual i's average Hamming "
        "distance to all other individuals in P(t), capturing how distinct/diverse the "
        "individual is relative to the rest of the population. The fitness bin f in {0,1,2,3} "
        "is the quartile bin of individual i's own fitness within the current population. "
        "Combining these gives 2 x 4 x 4 = 32 possible observation states, each mapping to a "
        "row of the Q-table Q(O,a)."
    )

    add_figure(
        doc, 4, "q2ga_action_selection.png",
        "Epsilon-greedy action selection and mapping to operator parameters",
        "Given the observation state O_t,i from Figure 3 and the linearly-decayed exploration "
        "rate epsilon(t) = epsilon_start + (epsilon_end - epsilon_start) * t/(T-1), the RL agent "
        "selects an action a in {1,...,5}. With probability epsilon(t) it explores by picking a "
        "uniformly random action; otherwise it exploits by choosing a = argmax_a Q(O_t,i, a). "
        "The chosen action a is then mapped to two operator parameters: n_rotate = "
        "max(1, round(a/5 * n_genes)), the number of genes (and corresponding qubits) that will "
        "be perturbed, and multiplier = a, an integer scale factor (1x to 5x) applied to the "
        "base rotation angle. n_rotate genes are then drawn at random (without replacement) "
        "from the n_genes real-valued genes of individual i, and these selections, together "
        "with the multiplier, are passed to the adaptive rotation gate U(t) described in "
        "Figure 5."
    )

    add_figure(
        doc, 5, "q2ga_rotation_gate.png",
        "Adaptive quantum rotation gate U(t)",
        "For each qubit j selected in Figure 4, the gate first computes sign(alpha_{i,j} * "
        "beta_{i,j}) (choosing +-1 at random if the product is zero). It then compares the "
        "collapsed bit x_{i,j} to the corresponding bit bb_j of the best-so-far solution: if "
        "x_{i,j} = 0 and bb_j = 1, the base rotation direction is +sign; if x_{i,j} = 1 and "
        "bb_j = 0, the base direction is -sign; if the bits already agree, the base direction "
        "is 0. With a small probability random_dir_frac (default 0.02), this direction is "
        "instead replaced by a uniformly random +-1, injecting stochastic exploration into the "
        "rotation direction independent of the best-so-far solution. The rotation angle is then "
        "theta_{i,j} = base_delta_theta * multiplier * direction, with base_delta_theta = "
        "0.05*pi by default and multiplier in {1,...,5} chosen by the RL agent (Figure 4). The "
        "standard 2D rotation matrix R(theta_{i,j}) is applied to (alpha_{i,j}, beta_{i,j}) and "
        "the result is renormalised so that alpha'^2 + beta'^2 = 1, producing the updated "
        "qubits that form row i of Q(t+1) -- this is the \"Update Q(t) with q-gate U(t)\" step "
        "in the overview chart."
    )

    add_figure(
        doc, 6, "q2ga_reward_update.png",
        "Reward calculation and Q-table update",
        "After the rotation gate updates individual i's qubits (Figure 5), a new bit string is "
        "sampled from the updated qubits and decoded into a new chromosome, which is evaluated "
        "to obtain new_fitness_i. The reward is computed as reward = (old_fitness_i - "
        "new_fitness_i) * 100 / |old_fitness_i|, so a reward greater than 0 corresponds to an "
        "improvement (since fitness is minimized) and a larger magnitude reflects a larger "
        "relative improvement (or, if negative, a relative deterioration). The visit counter "
        "counts(O_t,i, a) is incremented and the Q-table entry is updated using a fixed "
        "learning-rate update Q(O_t,i, a) <- Q(O_t,i, a) + lr * (reward - Q(O_t,i, a)), with "
        "lr = learning_rate (default 0.5) if provided, or 1/counts(O_t,i, a) otherwise (a "
        "running-average update). Finally, if new_fitness_i improves on the best fitness found "
        "so far, the best-so-far bit string and fitness are updated for use in the next "
        "generation's rotation gate (Figure 5). After all individuals are processed, the best "
        "individual of P(t) is saved (\"Save best individual of P(t)\") and the loop proceeds "
        "to generation t+1."
    )

    add_figure(
        doc, 7, "qga_vs_q2ga_rotation.png",
        "Comparison: QGA's fixed-angle rotation gate vs. Q2GA's adaptive rotation gate",
        "The QGA baseline applies a single fixed-angle rotation gate to every qubit of every "
        "individual in every generation: the rotation direction is determined purely by "
        "comparing each collapsed bit to the best-so-far genome's bits (no random-direction "
        "term), and the rotation angle is always DELTA_THETA = 0.05*pi, regardless of the "
        "individual's state or how the search is progressing. In contrast, Q2GA applies the "
        "rotation gate adaptively: only n_rotate genes' qubits (chosen per individual, based on "
        "the RL action) are rotated; the direction is usually toward the best-so-far genome but "
        "is occasionally randomized (probability random_dir_frac) for exploration; and the "
        "rotation magnitude is scaled by an RL-selected multiplier in {1,...,5}. Over many "
        "generations, the Q-table learns which (state, action) pairs -- i.e., which combinations "
        "of population-improvement trend, individual diversity, individual fitness rank, and "
        "rotation aggressiveness -- tend to yield the largest fitness improvements, allowing "
        "Q2GA to adapt its search intensity and direction in a way that the fixed-angle QGA "
        "cannot."
    )

    doc.add_page_break()
    add_heading(doc, "2. Algorithm Pseudocode", level=1)
    doc.add_paragraph(
        "The pseudocode below reflects the actual implementations in src/algorithms/ for the "
        "four algorithms compared in this study: the real-coded Genetic Algorithm (GA), "
        "Simulated Annealing (SA), the Quantum Genetic Algorithm (QGA, fixed-angle rotation), "
        "and the proposed RL-assisted Quantum Genetic Algorithm (Q2GA, adaptive rotation)."
    )

    add_pseudocode(doc, "2.1 Genetic Algorithm (GA)", [
        "Algorithm GA(pop_size, n_genes, T, pc, pm, eta=15, elitism=2)",
        "  Initialize population P(0): pop_size real-coded chromosomes in [0,1]^n_genes",
        "  Evaluate fitness(x) for all x in P(0)",
        "  best <- individual with minimum fitness in P(0)",
        "  for t = 0 to T-1 do",
        "    Sort P(t) by fitness; carry the top `elitism` individuals to P(t+1) unchanged",
        "    while |P(t+1)| < pop_size do",
        "      parent1 <- Tournament(P(t), k=3)   // pick best of 3 random individuals",
        "      parent2 <- Tournament(P(t), k=3)",
        "      with probability pc:",
        "        (child1, child2) <- SBX_Crossover(parent1, parent2, eta=15)",
        "      else:",
        "        (child1, child2) <- (parent1, parent2)",
        "      child1 <- Mutate(child1, pm); child2 <- Mutate(child2, pm)",
        "      clip child1, child2 to [0,1]^n_genes",
        "      add child1, child2 to P(t+1)",
        "    Evaluate fitness(x) for all x in P(t+1)",
        "    if min fitness in P(t+1) < fitness(best): best <- argmin individual",
        "    record best fitness for generation t",
        "  return best, history",
    ])

    add_pseudocode(doc, "2.2 Simulated Annealing (SA)", [
        "Algorithm SA(pop_size, n_genes, T, T0=10.0, alpha=0.97, sigma=0.15)",
        "  Initialize population P(0): pop_size real-coded chromosomes in [0,1]^n_genes",
        "  Evaluate fitness(x) for all x in P(0)",
        "  best <- individual with minimum fitness in P(0)",
        "  Temp <- T0",
        "  for t = 0 to T-1 do",
        "    for k = 1 to pop_size do   // pop_size perturbation trials per \"generation\"",
        "      x <- P(t)[k]",
        "      n_flip <- random integer in [1, max(2, n_genes // 5)]",
        "      x' <- copy of x with n_flip randomly chosen genes perturbed by",
        "            Gaussian noise (mean 0, std = sigma), clipped to [0,1]",
        "      if fitness(x') < fitness(x):",
        "        accept: x <- x'",
        "      else:",
        "        accept x <- x' with probability exp( -(fitness(x') - fitness(x)) / Temp )",
        "             (Metropolis criterion)",
        "      P(t)[k] <- x",
        "    P(t+1) <- P(t)",
        "    Temp <- Temp * alpha   // geometric cooling",
        "    if min fitness in P(t+1) < fitness(best): best <- argmin individual",
        "    record best fitness for generation t",
        "  return best, history",
    ])

    add_pseudocode(doc, "2.3 Quantum Genetic Algorithm (QGA, fixed-angle)", [
        "Algorithm QGA(pop_size, n_qubits, n_genes, T, DELTA_THETA = 0.05*pi)",
        "  Initialize Q(0): pop_size individuals, each n_qubits qubit pairs",
        "         (alpha_{i,j}, beta_{i,j}) = (1/sqrt(2), 1/sqrt(2)) for all i, j",
        "  best_bits <- None; best_fitness <- +infinity",
        "  for t = 0 to T-1 do",
        "    // -- Make P(t) by observing Q(t) --",
        "    for each individual i:",
        "      for each qubit j: sample x_{i,j} = 1 w.p. beta_{i,j}^2, else 0",
        "      b_i <- (x_{i,1}, ..., x_{i,n_qubits});  chrom_i <- bits_to_real(b_i)",
        "      fitness_i <- evaluate(chrom_i)",
        "      if fitness_i < best_fitness: best_bits <- b_i; best_fitness <- fitness_i",
        "    // -- Update Q(t) with fixed-angle rotation gate --",
        "    for each individual i:",
        "      for each qubit j = 1..n_qubits:",
        "        s <- sign(alpha_{i,j} * beta_{i,j})  (random +-1 if zero)",
        "        if x_{i,j} = 0 and best_bits_j = 1: direction <- +s",
        "        elif x_{i,j} = 1 and best_bits_j = 0: direction <- -s",
        "        else: direction <- 0",
        "        theta <- DELTA_THETA * direction",
        "        (alpha_{i,j}, beta_{i,j}) <- R(theta) * (alpha_{i,j}, beta_{i,j})",
        "        renormalize (alpha_{i,j}, beta_{i,j})",
        "    record best_fitness for generation t",
        "  return best_bits, best_fitness, history",
    ])

    add_pseudocode(doc, "2.4 RL-assisted Quantum Genetic Algorithm (Q2GA)", [
        "Algorithm Q2GA(pop_size, n_qubits, n_genes, T, N_ACTIONS=5,",
        "                base_delta_theta = 0.05*pi, random_dir_frac = 0.02,",
        "                epsilon_start, epsilon_end, learning_rate,",
        "                GEN_WINDOW = 5, HAM_BINS = 4, FIT_BINS = 4)",
        "  Initialize Q(0): pop_size individuals, each n_qubits qubit pairs",
        "         (alpha_{i,j}, beta_{i,j}) = (1/sqrt(2), 1/sqrt(2)) for all i, j",
        "  Initialize Q-table Q(O,a) <- 0 for all 32 states O and a in {1,...,5}",
        "  Initialize counts(O,a) <- 0",
        "  best_bits <- None; best_fitness <- +infinity",
        "  mean_fit_history <- empty list",
        "",
        "  for t = 0 to T-1 do",
        "    epsilon(t) <- epsilon_start + (epsilon_end - epsilon_start) * t/(T-1)",
        "",
        "    // -- Stage 1: Make P(t) by observing Q(t); evaluate --",
        "    for each individual i:",
        "      sample b_i from (alpha_i, beta_i);  chrom_i <- bits_to_real(b_i)",
        "      old_fitness_i <- evaluate(chrom_i)",
        "      if old_fitness_i < best_fitness: best_bits <- b_i; best_fitness <- old_fitness_i",
        "    mean_fit(t) <- mean(old_fitness_i over i)",
        "    g <- 1 if mean_fit(t) < average(mean_fit_history[-GEN_WINDOW:]) else 0",
        "    append mean_fit(t) to mean_fit_history",
        "",
        "    // -- Stage 2: build observation O_t,i = (g, h, f) for each individual --",
        "    for each individual i:",
        "      ham_i <- mean_{k != i} HammingDistance(b_i, b_k)",
        "      h_i <- quartile_bin(ham_i over all individuals; HAM_BINS bins)",
        "      f_i <- quartile_bin(old_fitness_i over all individuals; FIT_BINS bins)",
        "      O_i <- (g, h_i, f_i)",
        "",
        "    // -- Stages 3-5: per individual, select action, rotate, reward, learn --",
        "    for each individual i:",
        "      // Stage 3: epsilon-greedy action selection",
        "      if rand() < epsilon(t):",
        "        a <- random action in {1,...,N_ACTIONS}",
        "      else:",
        "        a <- argmax_{a'} Q(O_i, a')",
        "      n_rotate <- max(1, round(a / N_ACTIONS * n_genes))",
        "      multiplier <- a",
        "      selected_genes <- n_rotate genes sampled without replacement",
        "",
        "      // Stage 4: adaptive rotation gate U(t)",
        "      for each qubit j corresponding to a selected gene:",
        "        s <- sign(alpha_{i,j} * beta_{i,j})  (random +-1 if zero)",
        "        if x_{i,j} = 0 and best_bits_j = 1: direction <- +s",
        "        elif x_{i,j} = 1 and best_bits_j = 0: direction <- -s",
        "        else: direction <- 0",
        "        with probability random_dir_frac: direction <- random choice in {-1, +1}",
        "        theta <- base_delta_theta * multiplier * direction",
        "        (alpha_{i,j}, beta_{i,j}) <- R(theta) * (alpha_{i,j}, beta_{i,j})",
        "        renormalize (alpha_{i,j}, beta_{i,j})",
        "",
        "      // Stage 5: reward calculation and Q-table update",
        "      sample b_i' from updated (alpha_i, beta_i);  chrom_i' <- bits_to_real(b_i')",
        "      new_fitness_i <- evaluate(chrom_i')",
        "      reward <- (old_fitness_i - new_fitness_i) * 100 / |old_fitness_i|",
        "      counts(O_i, a) <- counts(O_i, a) + 1",
        "      lr <- learning_rate if learning_rate is not None else 1 / counts(O_i, a)",
        "      Q(O_i, a) <- Q(O_i, a) + lr * (reward - Q(O_i, a))",
        "      if new_fitness_i < best_fitness: best_bits <- b_i'; best_fitness <- new_fitness_i",
        "",
        "    Save best individual of P(t)",
        "    record best_fitness for generation t",
        "",
        "  return best_bits, best_fitness, history, Q",
    ])

    doc.save(OUT_PATH)
    print("Saved", OUT_PATH)


if __name__ == "__main__":
    main()
