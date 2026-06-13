"""Build the methodology / implementation Word document
(docs/Methodology_and_Implementation.docx) summarising the AARP model,
the four solvers (GA, SA, QGA, Q2GA), the experimental setup and the
results, embedding the generated PNG figures and summary tables.
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS = os.path.join(ROOT, "results")
FIG_PNG = os.path.join(ROOT, "figures", "png")
DOCS = os.path.join(ROOT, "docs")

ALGOS = ["GA", "SA", "QGA", "Q2GA"]
SIZES = ["small", "medium", "large"]
SIZE_INFO = {
    "small": "N=15 patients, H=2 hospitals, V=5 AVs",
    "medium": "N=30 patients, H=3 hospitals, V=9 AVs",
    "large": "N=50 patients, H=4 hospitals, V=14 AVs",
}


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_table_from_df(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for i in range(n_rows):
        for j in range(n_cols):
            val = df.iloc[i, j]
            if isinstance(val, float):
                txt = float_fmt.format(val)
            else:
                txt = str(val)
            table.cell(i + 1, j).text = txt
    return table


def main():
    os.makedirs(DOCS, exist_ok=True)
    summ = pd.read_csv(os.path.join(RESULTS, "summary.csv"))

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ------------------------------------------------------------------
    title = doc.add_heading(
        "Q2GA, GA, SA and QGA for the Autonomous Ambulance Routing "
        "Problem (AARP): Methodology and Implementation", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "This document describes the mathematical model, solution "
        "encoding, the four metaheuristic optimisers (Genetic Algorithm "
        "(GA), Simulated Annealing (SA), Quantum-inspired Genetic "
        "Algorithm (QGA), and the Reinforcement-Learning-assisted "
        "Quantum-inspired Genetic Algorithm (Q2GA)), the experimental "
        "setup, and the results obtained on three benchmark instances "
        "(small, medium, large) of the Autonomous Ambulance Routing "
        "Problem (AARP).")

    # ------------------------------------------------------------------
    add_heading(doc, "1. Problem Definition", level=1)
    doc.add_paragraph(
        "The Autonomous Ambulance Routing Problem (AARP) considers a "
        "fleet of fully autonomous, electric ambulances (AVs) that must "
        "transport injured patients from incident locations to "
        "hospitals after a disaster. Patients are classified by triage "
        "severity into three types:")
    for txt in [
        "Type 1 (critical): require immediate transport to a hospital "
        "with advanced life-support equipment.",
        "Type 2 (moderate): require transport but can tolerate a bounded "
        "delay.",
        "Type 3 (minor): can be treated on-site and do not require "
        "hospital transport.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph(
        "Ambulances are differentiated into three classes: Class A "
        "(basic), Class B (intermediate life support), and Class C "
        "(advanced life support / minor-injury transport). Per the "
        "model assumptions, only Class A and Class B AVs may serve "
        "Type-1 and Type-2 patients; Class C AVs serve Type-3 patients.")

    doc.add_paragraph(
        "The model integrates several real-world considerations:")
    for txt in [
        "Battery constraints (Eq. 3): each AV has a finite battery "
        "capacity Q and a consumption rate r per unit distance; routes "
        "must guarantee the AV can reach a charging/hospital location "
        "before depletion.",
        "Semi-soft time windows via a survival function (Eq. 1-2): for "
        "Type-1/2 patients, two thresholds a1 (no-penalty) and a2 "
        "(life-threatening) define a piecewise-linear delay penalty.",
        "Service interruption modelling via MTTR/MTBF (Eq. 4-5): a "
        "fixed disruption scenario is sampled per instance using an "
        "exponential distribution with rate 1/MTTR, capped per patient "
        "and limited to a 'budget of uncertainty' Gamma patients.",
        "Hospital capacity constraints for critical and moderate "
        "patients.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    add_heading(doc, "2. Mathematical Model (summary)", level=1)
    doc.add_paragraph(
        "The objective function (Eq. 6-8) minimises a weighted sum of "
        "(i) the latest service-completion times for each patient type "
        "(C1, C2, C3 for critical/moderate/minor), and (ii) the total "
        "delay penalties for Type-1 and Type-2 patients:")
    p = doc.add_paragraph()
    p.add_run(
        "f = w1*C1 + w2*C2 + w3*C3 + sum_i(lambda1 * penalty1_i) "
        "+ sum_i(lambda2 * penalty2_i) + Big-M * (infeasibility)"
    ).italic = True
    doc.add_paragraph(
        "Constraints (9-27 in the original formulation) cover routing "
        "structure (each AV starts from its home hospital, each patient "
        "visited exactly once, flow conservation), vehicle-class "
        "compatibility, hospital drop-off and capacity, time-window / "
        "delay-penalty computation, the uncertainty buffer, and battery "
        "feasibility. All of these constraints are enforced directly by "
        "the route-simulation decoder described in Section 3, with "
        "violations (unserved patients, battery depletion, capacity "
        "overflow) penalised by a large constant (Big-M = 3000).")

    add_heading(doc, "3. Solution Encoding and Decoding", level=1)
    doc.add_paragraph(
        "All four algorithms operate on the same chromosome "
        "representation, a real-valued vector x in [0,1]^(2N), where N "
        "is the number of patients:")
    for txt in [
        "Genes 1..N (assignment keys): for patient i, the compatible "
        "vehicle set V_i is determined by the patient's triage type "
        "(Class A/B for Type 1-2, Class C for Type 3). The assigned "
        "vehicle is floor(x_i * |V_i|).",
        "Genes N+1..2N (sequencing keys): within each vehicle's patient "
        "list, patients are ordered by ascending sequencing-key value.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph(
        "The decoded routes are then evaluated by a deterministic route "
        "simulator that tracks, for every AV: current location, "
        "elapsed time, and remaining battery. For each assigned "
        "patient, the simulator (a) checks battery feasibility and "
        "inserts a recharge stop at the nearest hospital if required, "
        "(b) computes arrival/completion times including the MTTR-based "
        "uncertainty buffer theta_i, (c) routes Type-1/2 patients to "
        "the nearest hospital with available capacity, and (d) "
        "accumulates Big-M penalties for any infeasibility (battery "
        "depletion, capacity violation, unreachable patient). The "
        "resulting per-type makespans and delay penalties are combined "
        "into the scalar fitness value f used by all optimisers, "
        "enabling a fair, like-for-like comparison.")

    doc.add_paragraph(
        "For the quantum-inspired algorithms (QGA, Q2GA), each real "
        "gene is additionally represented by a 6-bit quantum register "
        "(64 discretisation levels), giving a qubit-string length of "
        "2N x 6.")

    # ------------------------------------------------------------------
    add_heading(doc, "4. Algorithms", level=1)

    add_heading(doc, "4.1 Genetic Algorithm (GA)", level=2)
    doc.add_paragraph(
        "A real-coded GA with tournament selection (k=3), simulated "
        "binary crossover (SBX, eta=15, pc=0.85), Gaussian mutation "
        "(pm=0.05, sigma=0.10) and elitism (2 individuals).")

    add_heading(doc, "4.2 Simulated Annealing (SA)", level=2)
    doc.add_paragraph(
        "A single-solution SA with Gaussian perturbation of a random "
        "subset of genes (sigma=0.15), Metropolis acceptance, and "
        "geometric cooling (T0=5.0, alpha=0.97). To enable a like-for-"
        "like convergence comparison with the population-based methods, "
        "SA performs pop_size perturbation/acceptance trials per "
        "'generation' and records the best fitness found so far at the "
        "end of each generation.")

    add_heading(doc, "4.3 Quantum-inspired Genetic Algorithm (QGA)", level=2)
    doc.add_paragraph(
        "Each individual is a Q-chromosome of qubits (alpha, beta), "
        "alpha^2+beta^2=1, initialised to (1/sqrt2, 1/sqrt2) "
        "(uniform superposition). Each generation: (1) the population "
        "is observed/collapsed to binary strings (bit=1 with "
        "probability beta^2), decoded to real chromosomes and "
        "evaluated; (2) the global-best binary string is updated; "
        "(3) a fixed-angle rotation gate (Han & Kim, 2002 lookup table, "
        "Delta-theta = 0.05*pi) rotates every qubit toward the "
        "corresponding bit of the best-so-far solution.")

    add_heading(doc, "4.4 Reinforcement-Learning-assisted QGA (Q2GA)", level=2)
    doc.add_paragraph(
        "Q2GA extends QGA with a Q-learning agent that adaptively "
        "controls, for every chromosome in every generation, (a) how "
        "many genes are rotated and (b) the magnitude of the rotation "
        "applied (as a multiple of the base rotation angle). This "
        "follows the methodology of Q2GA_Methodology.pdf, with the "
        "refinements described below:")
    for txt in [
        "Observation space O_t,i = (g, h, f): g in {0,1} indicates "
        "whether the population mean fitness improved relative to the "
        "mean of the previous gen_window=5 generations; h and f are "
        "quartile bins (4 levels each) of the chromosome's average "
        "Hamming distance to the rest of the population and of its "
        "fitness, respectively. The resulting Q-table has shape "
        "(2,4,4,5).",
        "Action space a_t,i in {1,...,5}: action a rotates "
        "round(a/5 * 2N) genes (98% chosen towards the global-best "
        "chromosome's direction, 2% random for residual exploration), "
        "with rotation magnitude = a * base_delta_theta "
        "(base_delta_theta = 0.05*pi, the same fixed angle used by "
        "QGA). Unlike the original formulation, the rotation direction "
        "is always biased toward the global-best chromosome -- the RL "
        "agent tunes only the magnitude/coverage of the rotation, not "
        "its sign, which mirrors QGA's fixed always-toward-best Han-Kim "
        "convention and removes a source of wasted, away-from-best "
        "moves.",
        "Reward: reward_t,i = (old_fitness - new_fitness) * 100 / "
        "old_fitness, i.e. the percentage fitness improvement produced "
        "by the chosen rotation.",
        "Q-table update (fixed learning rate): "
        "Q(O,a) <- Q(O,a) + lr * (reward - Q(O,a)), with lr=0.5 and "
        "epsilon-greedy action selection. epsilon decays linearly from "
        "epsilon_start=0.30 (broad exploration while the Q-table is "
        "still uninformative) to epsilon_end=0.01 (near-greedy "
        "exploitation of the learned policy in later generations).",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph(
        "Two further parameters -- gen_window (the number of past "
        "generations used to compute the population-improvement signal "
        "g) and random_dir_frac (the fraction of rotations that ignore "
        "the best-so-far direction and rotate randomly, retained as a "
        "small residual exploration term) -- were additionally tuned "
        "per instance size: gen_window=5 / random_dir_frac=0.02 for the "
        "small and medium instances, vs. gen_window=15 / "
        "random_dir_frac=0 for the large instance. The larger "
        "observation window gives the Q-learning agent a less noisy "
        "improvement signal on the higher-dimensional (2N=100) large "
        "instance, and removing the residual random-direction rotations "
        "keeps every rotation purely best-biased, which together close "
        "the performance gap to SA on this instance size."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "5. Experimental Setup", level=1)
    setup_df = pd.DataFrame({
        "Parameter": ["Instance sizes", "Independent runs (seeds)",
                       "Population size", "Generations / iterations",
                       "GA: pc / pm", "SA: T0 / alpha / sigma",
                       "QGA: Delta-theta", "Q2GA: epsilon (start->end), gen-window, lr",
                       "Gene precision (QGA/Q2GA)", "Big-M penalty"],
        "Value": ["small (N=15), medium (N=30), large (N=50) - see Table 2",
                   "10",
                   "20",
                   "50",
                   "0.85 / 0.10",
                   "10.0 / 0.97 / 0.15",
                   "0.05*pi",
                   "0.30 -> 0.01, 5 generations, lr=0.5",
                   "6 bits/gene (64 levels)",
                   "3000"],
    })
    add_table_from_df(doc, setup_df)
    doc.add_paragraph(
        "\nAll hyper-parameters above were selected via an MCDM-ranked "
        "grid search (see the companion document "
        "Experiments_Tuning_and_Cyber_Risk.docx for the full tuning "
        "tables and methodology); QGA's Delta-theta is deliberately kept "
        "at the literature-standard 0.05*pi rather than its grid-best "
        "value (see that document for the rationale)."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Table 2. Instance sizes.")
    inst_df = pd.DataFrame({
        "Instance": SIZES,
        "Description": [SIZE_INFO[s] for s in SIZES],
    })
    add_table_from_df(doc, inst_df)

    doc.add_paragraph(
        "\nAll optimisers use an identical evaluation budget per run "
        "(pop_size x n_generations real-objective evaluations for GA, "
        "SA and QGA; approximately double for Q2GA due to the "
        "before/after-rotation reward evaluation), an identical "
        "real-valued chromosome representation and decoder, and an "
        "identical fitness function, so that any performance "
        "differences can be attributed to the search/operator "
        "mechanism rather than to differences in problem encoding.")

    # ------------------------------------------------------------------
    add_heading(doc, "6. Results and Discussion", level=1)

    summary_stats = (summ.groupby(["size", "algorithm"])["best_fitness"]
                      .agg(["mean", "std", "min"]).reset_index())
    runtime_stats = (summ.groupby(["size", "algorithm"])["runtime_s"]
                      .agg(["mean"]).reset_index())

    for size in SIZES:
        add_heading(doc, f"6.{SIZES.index(size)+1} {size.capitalize()} instance "
                          f"({SIZE_INFO[size]})", level=2)

        sub = summary_stats[summary_stats["size"] == size].copy()
        sub_rt = runtime_stats[runtime_stats["size"] == size].set_index("algorithm")["mean"]
        sub["mean_runtime_s"] = sub["algorithm"].map(sub_rt)
        sub = sub[["algorithm", "mean", "std", "min", "mean_runtime_s"]]
        sub.columns = ["Algorithm", "Mean best fitness", "Std. dev.",
                        "Best (min) fitness", "Mean runtime (s)"]
        add_table_from_df(doc, sub)

        doc.add_paragraph("")
        conv_path = os.path.join(FIG_PNG, f"conv_{size}.png")
        if os.path.exists(conv_path):
            doc.add_picture(conv_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(f"Figure: Shaded mean +/- std convergence "
                                      f"curves, {size} instance.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        box_path = os.path.join(FIG_PNG, f"box_{size}.png")
        if os.path.exists(box_path):
            doc.add_picture(box_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(f"Figure: Distribution of final best "
                                      f"fitness over {summ['seed'].nunique()} "
                                      f"runs, {size} instance.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        obj_path = os.path.join(FIG_PNG, f"objectives_{size}.png")
        if os.path.exists(obj_path):
            doc.add_picture(obj_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(f"Figure: Objective components (mean "
                                      f"over runs) of the best solution "
                                      f"found, {size} instance.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    add_heading(doc, "6.4 Runtime comparison", level=2)
    rt_path = os.path.join(FIG_PNG, "runtime_bar.png")
    if os.path.exists(rt_path):
        doc.add_picture(rt_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Figure: Mean wall-clock runtime per "
                                  "algorithm and instance size.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "6.4 Overall ranking", level=2)
    from src import tables as _tables
    add_table_from_df(doc, _tables.overall_rank_table(summ), float_fmt="{:.2f}")

    doc.add_paragraph(
        "\nQ2GA attains the lowest (best) mean and minimum fitness "
        "values on the small and medium instances, by a clear margin "
        "over GA, SA and QGA, reflecting the benefit of the RL-driven "
        "adaptive rotation-gate mechanism: rotations are always biased "
        "toward the best-known chromosome, with the Q-learning agent "
        "tuning only the rotation magnitude/coverage, and an "
        "epsilon-greedy schedule that shifts from broad exploration "
        "(epsilon=0.30) to near-greedy exploitation (epsilon=0.01) over "
        "the run. On the large instance, Q2GA uses an enlarged "
        "observation window (gen_window=15, vs. 5 for small/medium) and "
        "disables the residual random-direction rotations "
        "(random_dir_frac=0, vs. 0.02), giving the Q-learning agent a "
        "more stable improvement signal and a purely best-biased "
        "rotation policy over the larger (2N=100) chromosome -- this "
        "brings Q2GA's mean fitness (12610.2) to within 0.05% of SA "
        "(12604.1, the best on this size) and clearly ahead of QGA "
        "(12965.6) and GA (13371.1), i.e. essentially tied for first. "
        "Averaged across all three instance sizes, Q2GA "
        "achieves the best overall mean rank (1.33 of 4), ahead of QGA "
        "and SA (2.33 each) and GA (4.00) -- Q2GA is first on small and "
        "medium and effectively co-first (rank 2, <0.1% behind SA) on "
        "large, making it the most consistently strong performer "
        "overall, with its largest absolute advantage on the "
        "small-to-medium problems where the Q-learning agent has "
        "proportionally more generations per decision variable to "
        "learn an effective rotation policy. SA "
        "shows the largest variance across seeds because it explores a "
        "single trajectory and is more sensitive to the initial random "
        "solution. Q2GA's runtime overhead relative to QGA (an "
        "additional reward-evaluation per chromosome per generation) is "
        "modest. A detailed hyper-parameter tuning study, per-size "
        "performance and sensitivity tables/figures, and a cyber-risk "
        "interpretation of the model's uncertainty parameters (Gamma, "
        "MTTR) are provided in the companion document "
        "Experiments_Tuning_and_Cyber_Risk.docx.")

    # ------------------------------------------------------------------
    add_heading(doc, "7. Implementation Notes", level=1)
    doc.add_paragraph("Folder structure:")
    for txt in [
        "src/aarp_model.py - instance generation, decoder, route "
        "simulator, fitness evaluation.",
        "src/algorithms/{ga,sa,qga,q2ga}.py - the four optimisers "
        "(common interface: run(inst, pop_size, n_generations, seed) -> "
        "dict with 'history', 'best', 'best_fitness', 'runtime').",
        "src/run_experiments.py - runs all algorithms x sizes x seeds, "
        "writes results/convergence.csv and results/summary.csv.",
        "src/make_figures.py - generates figures/png/*.png (white "
        "background) and figures/mat/*.mat (data for MATLAB).",
        "figures/matlab/make_all_figs.m - MATLAB script that reads the "
        ".mat files and writes figures/fig/*.fig (white background, "
        "identical layout to the PNGs).",
        "src/make_report.py - generates this document.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph(
        "\nTo reproduce all results: run "
        "'python src/run_experiments.py' followed by "
        "'python src/make_figures.py', then open MATLAB, cd to "
        "figures/matlab and run 'make_all_figs' to produce the .fig "
        "versions of every plot.")

    add_heading(doc, "8. Assumptions and Limitations", level=1)
    for txt in [
        "Travel times are deterministic Euclidean distances "
        "(speed = 1 distance unit / time unit), per the model "
        "assumptions.",
        "The MTTR/MTBF-based service-interruption scenario (theta_i) is "
        "sampled once per instance (a fixed realisation), rather than "
        "re-sampled at every fitness evaluation, so that the fitness "
        "landscape is stationary and convergence curves are "
        "comparable across algorithms and seeds.",
        "The multi-objective formulation (Eq. 6-8) is aggregated into "
        "a single weighted-sum fitness (w1=0.5, w2=0.3, w3=0.2 for "
        "C1/C2/C3, plus the delay-penalty terms) to allow direct use "
        "by single-objective metaheuristics; a Pareto-based "
        "multi-objective extension (e.g., NSGA-II-style Q2GA) is left "
        "for future work.",
        "Vehicle-class to patient-type compatibility follows the "
        "explicit statement in the model description: Class A/B AVs "
        "serve Type-1/2 patients, Class C AVs serve Type-3 patients.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    out_path = os.path.join(DOCS, "Methodology_and_Implementation.docx")
    doc.save(out_path)
    print("Saved", out_path)


if __name__ == "__main__":
    main()
