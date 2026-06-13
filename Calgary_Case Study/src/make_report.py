"""Build the Calgary case-study Word document (docs/Calgary_Case_Study.docx).

Explains the case-study motivation, real-world data sources, the use of a
real road-network travel-distance matrix (instead of Euclidean distance),
instance characteristics, algorithm-comparison results, cyber-risk /
availability sensitivity results, and the OpenStreetMap maps.
"""

import os
import sys
import json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ROOT_AARP = os.path.join(os.path.dirname(ROOT), "Q2GA_AARP")
sys.path.insert(0, ROOT_AARP)
sys.path.insert(0, ROOT)

RESULTS = os.path.join(ROOT, "results")
FIG_PNG = os.path.join(ROOT, "figures", "png")
MAPS = os.path.join(ROOT, "maps")
DOCS = os.path.join(ROOT, "docs")

from src.calgary_instance import build_instance, HOSPITALS, PATIENTS  # noqa: E402

ALGOS = ["GA", "SA", "QGA", "Q2GA"]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_table_from_df(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for i in range(n_rows):
        for j in range(n_cols):
            val = df.iloc[i, j]
            if isinstance(val, (float, np.floating)):
                txt = float_fmt.format(val)
            else:
                txt = str(val)
            table.cell(i + 1, j).text = txt
    return table


def add_fig(doc, path_dir, filename, width=6.0, caption=None):
    path = os.path.join(path_dir, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        if caption:
            doc.add_paragraph(caption)
    else:
        doc.add_paragraph(f"[missing figure: {filename}]")


def main():
    os.makedirs(DOCS, exist_ok=True)
    inst = build_instance(seed=42)
    summ = pd.read_csv(os.path.join(RESULTS, "summary.csv"))
    with open(os.path.join(RESULTS, "best_routes.json")) as f:
        best = json.load(f)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Calgary Case Study: Cyber-Resilient Autonomous Ambulance Routing\n"
                     "with Real Road-Network Travel Distances", level=0)
    doc.add_paragraph(
        "This document applies the Genetic-Algorithm (GA), Simulated-Annealing "
        "(SA), Quantum-inspired Genetic Algorithm (QGA) and the proposed "
        "Quantum-inspired Genetic Algorithm with Adaptive Annealing and Random "
        "Perturbation (Q2GA) -- as described in Methodology_and_Implementation.docx "
        "and Experiments_Tuning_and_Cyber_Risk.docx -- to a realistic case study "
        "of autonomous-ambulance dispatch in Calgary, Canada. Unlike the Ottawa "
        "case study (which uses straight-line / Euclidean distances on a local "
        "projection), this Calgary case study replaces all travel distances with "
        "real road-network driving distances obtained from OpenStreetMap, so "
        "that travel costs reflect the actual street network (arterials, "
        "ring roads, river crossings, etc.) rather than \"as the crow flies\" "
        "geometry. This document covers the real-world data used to build the "
        "instance, the road-network distance methodology, the resulting "
        "instance characteristics, the algorithm-comparison results, a "
        "cyber-risk / system-availability sensitivity analysis, and "
        "OpenStreetMap visualisations of the locations and the best routing "
        "plan found."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "1. Case-Study Motivation and Real-World Data", level=1)
    doc.add_paragraph(
        "Autonomous electric ambulances dispatched from hospital depots are "
        "increasingly proposed as a way to reduce emergency-response times in "
        "mid-sized cities. Such fleets depend on continuous vehicle-to-infrastructure "
        "communication for dispatch, re-routing and telemetry; a cyber attack or "
        "communication outage that degrades this link directly threatens patient "
        "outcomes. The Calgary case study instantiates the AARP model with real "
        "hospital locations and real residential-neighbourhood centroids, and -- "
        "going a step further than the Ottawa case study -- evaluates every "
        "candidate route using the actual Calgary road network, so that the "
        "cyber-risk / availability analysis from the main study can be related "
        "to a concrete, geographically and topologically realistic deployment."
    )

    add_heading(doc, "1.1 Hospitals (vehicle depots)", level=2)
    hosp_rows = []
    for h_idx, h in enumerate(HOSPITALS):
        hosp_rows.append({
            "ID": f"H{h_idx}",
            "Hospital": h["name"],
            "Latitude": h["lat"],
            "Longitude": h["lon"],
        })
    add_table_from_df(doc, pd.DataFrame(hosp_rows), float_fmt="{:.4f}")
    doc.add_paragraph(
        "These three hospitals are real Calgary hospital sites (Foothills "
        "Medical Centre in the northwest, Peter Lougheed Centre in the "
        "northeast, and Rockyview General Hospital in the south), providing a "
        "realistic multi-depot configuration spanning the city. Vehicles are "
        "assigned home depots and triage-related capacity limits "
        "(hosp_cap_critical, hosp_cap_moderate) at each hospital, exactly as in "
        "the main AARP model."
    )

    add_heading(doc, "1.2 Patient locations (residential neighbourhoods)", level=2)
    doc.add_paragraph(
        "25 patient locations were placed at the approximate centroids of real "
        "Calgary residential neighbourhoods, spanning the northwest (Tuscany, "
        "Royal Oak, Bowness, Montgomery, Aspen Woods, Springbank Hill, West "
        "Springs, Signal Hill, Killarney), the north (Evanston, Panorama Hills, "
        "Saddle Ridge, Taradale), the east (Marlborough, Forest Lawn), the "
        "inner city (Inglewood, Ogden), and the south/southeast (Riverbend, "
        "McKenzie Towne, Mahogany, Auburn Bay, Cranston, Sundance, New "
        "Brighton, Copperfield). Each location was independently assigned a "
        "triage type (1 = critical, 2 = moderate, 3 = minor) using the same "
        "probability mix as the main study's medium instance "
        "(p = [0.25, 0.35, 0.40])."
    )
    pat_rows = []
    for i, p in enumerate(PATIENTS):
        pat_rows.append({
            "ID": f"P{i}",
            "Neighbourhood": p["name"],
            "Latitude": p["lat"],
            "Longitude": p["lon"],
            "Type": int(inst.patient_type[i]),
        })
    add_table_from_df(doc, pd.DataFrame(pat_rows), float_fmt="{:.4f}")

    type_counts = {t: int((inst.patient_type == t).sum()) for t in (1, 2, 3)}
    doc.add_paragraph(
        f"Resulting triage mix for this instance (seed=42): "
        f"{type_counts[1]} critical, {type_counts[2]} moderate, "
        f"{type_counts[3]} minor patients."
    )
    add_fig(doc, FIG_PNG, "patient_demographics.png", width=4.5)

    # ------------------------------------------------------------------
    add_heading(doc, "2. Road-Network Travel Distances (No Euclidean Distance)", level=1)
    doc.add_paragraph(
        "The AARP simulator (src/aarp_model.py) was extended with an optional "
        "`dist_matrix` field on the problem `Instance`. When present, every "
        "travel-distance lookup between two locations (hospital-to-patient, "
        "patient-to-hospital, hospital-to-hospital, including the battery-"
        "recharge and hospital drop-off legs) is taken from this matrix "
        "instead of being computed as a straight-line (Euclidean) distance "
        "between coordinates. The Calgary case study populates `dist_matrix` "
        "as follows:"
    )
    steps = pd.DataFrame([
        {"Step": "1", "Description": "Download the OpenStreetMap drive-network graph for the "
                                       "bounding box covering all 3 hospitals and 25 patient "
                                       "locations (with a 0.05-degree padding) using `osmnx`."},
        {"Step": "2", "Description": "Snap each of the 28 locations (3 hospitals + 25 patients) "
                                       "to its nearest road-network node."},
        {"Step": "3", "Description": "For each location, run Dijkstra's shortest-path algorithm "
                                       "over the road network (edge weight = road segment length "
                                       "in metres) to every other location."},
        {"Step": "4", "Description": "Assemble the resulting 28x28 matrix of shortest road "
                                       "distances (km) as `dist_matrix`, and cache it "
                                       "(results/dist_matrix.npy) together with the downloaded "
                                       "road graph (results/road_network.graphml)."},
    ])
    add_table_from_df(doc, steps)

    doc.add_paragraph(
        "Latitude/longitude pairs are still converted to a local Cartesian "
        "(km) frame via an equirectangular projection centred on Calgary "
        "(51.0447 N, -114.0719 W) for plotting purposes (instance-layout "
        "figures, etc.), but this projected coordinate distance is no longer "
        "used by the route simulator -- only `dist_matrix` (real road-network "
        "driving distance) is used to compute travel times, battery "
        "consumption, and feasibility checks. The same real-world unit "
        "mapping as the Ottawa case study is retained:"
    )
    unit_rows = pd.DataFrame([
        {"Quantity": "Distance unit", "Value": "1 km (road-network distance, not straight-line)"},
        {"Quantity": "Time unit", "Value": "1 minute"},
        {"Quantity": "Vehicle speed", "Value": "0.5 km/min (30 km/h average urban speed)"},
        {"Quantity": "Battery range Q", "Value": "300 km"},
        {"Quantity": "Battery consumption r", "Value": "1 unit / km"},
    ])
    add_table_from_df(doc, unit_rows)

    max_d = float(inst.dist_matrix.max())
    doc.add_paragraph(
        f"The downloaded Calgary drive network spans the full extent of the "
        f"28 locations (the patient locations alone span roughly "
        f"{inst.coords_patients[:, 0].max() - inst.coords_patients[:, 0].min():.1f} km "
        f"east-west by "
        f"{inst.coords_patients[:, 1].max() - inst.coords_patients[:, 1].min():.1f} km "
        f"north-south in straight-line terms); the largest pairwise "
        f"road-network distance in `dist_matrix` is {max_d:.1f} km, "
        f"reflecting the additional travel required to follow Calgary's "
        f"arterial road and ring-road layout rather than a direct line."
    )
    add_fig(doc, FIG_PNG, "instance_layout_xy.png", width=5.5,
            caption="Figure 1. Instance layout in the local (km) coordinate frame "
                    "(for visualisation only -- travel costs use road-network "
                    "distances, not these straight-line coordinates).")
    add_fig(doc, MAPS, "static_map.png", width=5.5,
            caption="Figure 2. Hospital and patient locations shown on an "
                    "OpenStreetMap basemap of Calgary.")

    # ------------------------------------------------------------------
    add_heading(doc, "3. Instance Characteristics", level=1)
    doc.add_paragraph(
        f"The Calgary instance has N={inst.n_patients} patients, "
        f"H={inst.n_hospitals} hospitals and M={inst.n_vehicles} vehicles -- "
        "between the main study's small (N=15) and medium (N=30) instances, "
        "matching the Ottawa case study's scale for comparability. The tuned "
        "hyper-parameters from the medium-instance configuration "
        "(Experiments_Tuning_and_Cyber_Risk.docx, Section 1) are reused for "
        "consistency:"
    )
    tuned_rows = pd.DataFrame([
        {"Algorithm": "GA", "Tuned parameters": "pc=0.85, pm=0.10"},
        {"Algorithm": "SA", "Tuned parameters": "T0=10.0, alpha=0.97"},
        {"Algorithm": "QGA", "Tuned parameters": "delta_theta=0.05*pi"},
        {"Algorithm": "Q2GA", "Tuned parameters": "epsilon_start=0.30, base_delta_theta=0.05*pi"},
    ])
    add_table_from_df(doc, tuned_rows)

    doc.add_paragraph("Service and drop-off time distributions for this instance:")
    add_fig(doc, FIG_PNG, "service_dropoff_times.png", width=6.0)

    doc.add_paragraph("Semi-soft time-window thresholds (a1, a2) by triage type:")
    add_fig(doc, FIG_PNG, "time_window_thresholds.png", width=5.0)
    add_fig(doc, FIG_PNG, "penalty_function_shape.png", width=5.5,
            caption="Figure 3. Delay-penalty function for the realised a1/a2 "
                    "thresholds in this instance.")

    add_fig(doc, FIG_PNG, "goal_function_weights.png", width=4.5,
            caption="Figure 4. Weighted-sum goal-function structure "
                    "(unchanged from the main study).")

    # ------------------------------------------------------------------
    add_heading(doc, "4. Algorithm-Comparison Results", level=1)
    doc.add_paragraph(
        f"All four algorithms were run with POP_SIZE=20, N_GENERATIONS=50, and "
        f"10 independent seeds, using the tuned hyper-parameters above, with "
        f"the route simulator evaluating travel costs via the real "
        f"road-network `dist_matrix`. Final best-fitness statistics:"
    )
    stats = summ.groupby("algorithm")["best_fitness"].agg(
        mean="mean", std="std", min="min", max="max").reset_index()
    stats = stats.set_index("algorithm").loc[ALGOS].reset_index()
    stats.columns = ["Algorithm", "Mean", "Std", "Min", "Max"]
    add_table_from_df(doc, stats, float_fmt="{:.2f}")

    runtime = summ.groupby("algorithm")["runtime_s"].mean().loc[ALGOS]
    best_overall = float(summ["best_fitness"].min())
    doc.add_paragraph(
        f"Q2GA achieves the lowest (best) mean final fitness "
        f"({stats.set_index('Algorithm').loc['Q2GA', 'Mean']:.2f}) with a "
        f"smaller spread (std = {stats.set_index('Algorithm').loc['Q2GA', 'Std']:.2f}) "
        f"than the non-RL algorithms, consistent with its overall ranking in "
        f"the main study and in the Ottawa case study. The single best solution "
        f"found across all 40 runs has fitness {best_overall:.2f} (Q2GA), and "
        f"is used as the routing plan visualised in Section 6. Mean runtimes "
        f"per 50-generation run were GA={runtime['GA']:.2f}s, "
        f"SA={runtime['SA']:.2f}s, QGA={runtime['QGA']:.2f}s, "
        f"Q2GA={runtime['Q2GA']:.2f}s -- Q2GA's higher cost reflects its "
        f"additional adaptive-annealing and random-perturbation steps."
    )

    add_fig(doc, FIG_PNG, "convergence.png", width=6.0,
            caption="Figure 5. Mean +/- standard-deviation convergence curves "
                    "(10 seeds).")
    add_fig(doc, FIG_PNG, "boxplot.png", width=5.0,
            caption="Figure 6. Distribution of final best fitness across seeds.")
    add_fig(doc, FIG_PNG, "objectives.png", width=6.0,
            caption="Figure 7. Mean objective-component breakdown of the best "
                    "solution found by each algorithm.")
    add_fig(doc, FIG_PNG, "runtime_bar.png", width=5.0,
            caption="Figure 8. Mean wall-clock runtime per algorithm.")

    comp = best["objective_components"]
    doc.add_paragraph(
        f"For the best Q2GA solution (fitness={best['best_fitness']:.2f}), the "
        f"objective decomposes as C1={comp['C1']:.2f} (critical completion), "
        f"C2={comp['C2']:.2f} (moderate completion), C3={comp['C3']:.2f} "
        f"(minor completion), penalty1={comp['penalty1']:.2f} (critical delay), "
        f"penalty2={comp['penalty2']:.2f} (moderate delay), with "
        f"{int(comp['n_unserved'])} unserved patients. Because completion times "
        f"(C1-C3) and delay penalties are now driven by real road-network "
        f"travel times rather than straight-line travel times, they are not "
        f"directly comparable in absolute terms to the Ottawa case study's "
        f"figures -- the Calgary values reflect the additional distance "
        f"travelled along actual streets. A residual infeasibility term of "
        f"{comp['infeasibility']:.0f} (= one BIG_M=3000 unit) remains, "
        f"corresponding to a single soft constraint (e.g. a hospital "
        f"triage-capacity limit) being exceeded by one unit -- a pattern also "
        f"observed at this problem scale in the main study and the Ottawa "
        f"case study, and a candidate for further tuning or an additional "
        f"vehicle in future work."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "5. Cyber-Risk and System-Availability Sensitivity", level=1)
    doc.add_paragraph(
        "Following the cyber-risk framing of Experiments_Tuning_and_Cyber_Risk.docx "
        "(Section 4), service disruptions caused by a cyber attack or "
        "communication outage are modelled via two parameters: the disruption "
        "budget Gamma (the number / fraction of patients whose dispatch is "
        "affected) and the mean time to repair (MTTR), which together with a "
        "fixed mean time between failures (MTBF=120 minutes) determine the "
        "system availability = MTBF / (MTBF + MTTR). For this case study, "
        "Q2GA (the recommended algorithm) was swept over Gamma in "
        "{0, 0.1, 0.2, 0.3, 0.4, 0.5} x N and MTTR in "
        "{2, 5, 10, 15, 20, 30} minutes, each with 5 seeds and 30 generations "
        "(a reduced budget relative to Section 4, for speed), again using the "
        "real road-network `dist_matrix` for all travel-cost evaluations."
    )

    add_fig(doc, FIG_PNG, "qos_vs_gamma.png", width=6.5,
            caption="Figure 9. Quality-of-service metrics vs. disruption "
                    "budget Gamma (Calgary case study, Q2GA, 5 seeds).")
    add_fig(doc, FIG_PNG, "qos_vs_mttr.png", width=6.0,
            caption="Figure 10. Quality-of-service metrics vs. MTTR, with the "
                    "corresponding system-availability scale (Calgary case "
                    "study, Q2GA, 5 seeds).")
    add_fig(doc, FIG_PNG, "availability_surface.png", width=5.5,
            caption="Figure 11. System availability as a function of MTTR and "
                    "MTBF, with the Calgary case study's operating point marked.")
    add_fig(doc, FIG_PNG, "disruption_scenario.png", width=6.0,
            caption="Figure 12. Realised per-patient disruption buffers theta_i "
                    f"for the default scenario (Gamma={inst.gamma_budget}, "
                    f"MTTR={inst.mttr:.0f}).")
    add_fig(doc, FIG_PNG, "cyber_risk_degradation_summary.png", width=5.5,
            caption="Figure 13. Summary of quality-of-service degradation under "
                    "two worsening cyber-risk scenarios.")

    doc.add_paragraph(
        "As in the main study and the Ottawa case study, increasing the MTTR "
        "(i.e. reducing system availability) increases the mean total delay "
        "penalty, while widening the disruption budget Gamma primarily affects "
        "which patients are delayed rather than producing a large change in "
        "overall fitness, since Q2GA's adaptive-annealing and "
        "random-perturbation mechanisms continue to find feasible re-routings "
        "around the affected patients for this instance size -- now over the "
        "real Calgary road network. The Calgary case study's default operating "
        f"point (Gamma={inst.gamma_budget} of {inst.n_patients} patients, "
        f"MTTR={inst.mttr:.0f} min, availability="
        f"{120.0/(120.0+inst.mttr):.2f}) sits in a region of high availability "
        "(>0.9), consistent with a well-maintained communication infrastructure; "
        "the sensitivity curves indicate that quality of service degrades "
        "noticeably once MTTR approaches 20-30 minutes (availability < 0.85), "
        "highlighting the importance of rapid incident response for "
        "cyber-resilience in this deployment."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "6. Best Routing Plan on the Calgary Map", level=1)
    doc.add_paragraph(
        f"The best solution found (Q2GA, fitness={best['best_fitness']:.2f}, "
        f"{int(comp['n_unserved'])} unserved patients) assigns each of the "
        f"{inst.n_vehicles} vehicles a route starting and ending at its home "
        f"hospital, with travel costs evaluated over the real Calgary road "
        f"network. The resulting routes are visualised on an OpenStreetMap "
        f"basemap below as straight lines between consecutive stops "
        f"(for clarity); the underlying fitness evaluation uses the actual "
        f"road-network distance for each leg, not the straight-line distance "
        f"shown. An interactive version with clickable markers and routes is "
        f"provided in maps/interactive_map.html and maps/route_map.html."
    )

    route_rows = []
    for v_idx_str, seq in best["routes"].items():
        v_idx = int(v_idx_str)
        home = int(inst.vehicle_home[v_idx])
        names = ", ".join(f"P{p} ({PATIENTS[p]['name']})" for p in seq)
        route_rows.append({
            "Vehicle": v_idx,
            "Home hospital": f"H{home}: {HOSPITALS[home]['name']}",
            "Stops (in order)": names if names else "(none)",
        })
    add_table_from_df(doc, pd.DataFrame(route_rows))

    add_fig(doc, MAPS, "route_map.png", width=6.5,
            caption="Figure 14. Best Q2GA routing plan overlaid on an "
                    "OpenStreetMap basemap of Calgary (straight lines shown "
                    "for clarity; costs use road-network distances).")

    # ------------------------------------------------------------------
    add_heading(doc, "7. Conclusions", level=1)
    doc.add_paragraph(
        "The Calgary case study extends the Ottawa case study's geographically "
        "grounded validation by replacing Euclidean (straight-line) travel "
        "distances with real road-network driving distances computed from "
        "OpenStreetMap data (via `osmnx` / `networkx`), through a single "
        "optional `dist_matrix` field added to the AARP `Instance`, with no "
        "other changes to the underlying simulator or solvers. On this more "
        "realistic instance, Q2GA again achieves the best mean and best-case "
        "objective values among GA, SA, QGA and Q2GA, at a modest additional "
        "computational cost, confirming that the RL-assisted adaptive rotation "
        "mechanism remains effective when travel costs follow the actual "
        "street network rather than idealised straight lines. The cyber-risk / "
        "availability sensitivity analysis again shows that the system "
        "tolerates moderate disruption budgets well, but that quality of "
        "service degrades as MTTR grows and availability falls below roughly "
        "0.85-0.90, underscoring the value of fast incident-response (low "
        "MTTR) for maintaining ambulance-dispatch quality of service under "
        "cyber-risk conditions -- now demonstrated on a real, road-network-"
        "constrained deployment."
    )

    out_path = os.path.join(DOCS, "Calgary_Case_Study.docx")
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    main()
